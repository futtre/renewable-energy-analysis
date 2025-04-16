from docx import Document
from fpdf import FPDF

# Create first DOCX file - Solar Project Agreement
doc1 = Document()
doc1.add_heading('POWER PURCHASE AGREEMENT: DESERT SUN SOLAR PROJECT', 0)
doc1.add_paragraph('This Power Purchase Agreement ("Agreement") is made and entered into on March 15, 2024')
doc1.add_paragraph('BETWEEN:')
doc1.add_paragraph('SunPower Development LLC ("Developer"), a renewable energy development company with its principal office at 1000 Solar Drive, Phoenix, AZ 85001')
doc1.add_paragraph('AND')
doc1.add_paragraph('Western Grid Utilities ("Offtaker"), a public utility company')

doc1.add_heading('PROJECT DETAILS', level=1)
doc1.add_paragraph('Project Name: Desert Sun Solar Project')
doc1.add_paragraph('Location: Mojave Desert, San Bernardino County, California')
doc1.add_paragraph('Technology: Utility-Scale Solar PV with Single-Axis Tracking')
doc1.add_paragraph('Capacity: 250 MW (AC)')

doc1.add_heading('KEY DATES', level=1)
doc1.add_paragraph('Notice to Proceed (NTP): Expected July 1, 2024')
doc1.add_paragraph('Construction Start: August 15, 2024')
doc1.add_paragraph('Commercial Operation Date (COD): Target December 31, 2025')

doc1.add_heading('PROJECT PARTICIPANTS', level=1)
doc1.add_paragraph('Developer: SunPower Development LLC')
doc1.add_paragraph('EPC Contractor: BuildRight Solar Construction Inc.')
doc1.add_paragraph('Landowner: Desert Holdings LLC')
doc1.add_paragraph('Offtaker: Western Grid Utilities')
doc1.save('test_files/solar_project_agreement.docx')

# Create second DOCX file - Wind Project Technical Specs
doc2 = Document()
doc2.add_heading('TECHNICAL SPECIFICATIONS: PRAIRIE WIND PROJECT', 0)
doc2.add_paragraph('PROJECT OVERVIEW')
doc2.add_paragraph('The Prairie Wind Project is a utility-scale wind energy facility being developed by WindCo Energy Partners. Located in central Iowa, this project represents a significant addition to the state\'s renewable energy portfolio.')

doc2.add_heading('PROJECT DETAILS', level=1)
doc2.add_paragraph('Project Name: Prairie Wind Project')
doc2.add_paragraph('Developer: WindCo Energy Partners')
doc2.add_paragraph('Location: Greene County, Iowa')
doc2.add_paragraph('Total Capacity: 175 MW')

doc2.add_heading('DEVELOPMENT TIMELINE', level=1)
doc2.add_paragraph('Contract Execution: February 1, 2024')
doc2.add_paragraph('Expected Construction Start: May 15, 2024')
doc2.add_paragraph('Target Commercial Operation: June 30, 2025')

doc2.add_heading('KEY STAKEHOLDERS', level=1)
doc2.add_paragraph('Project Owner: WindCo Energy Partners')
doc2.add_paragraph('Turbine Supplier: VestasWind Systems')
doc2.add_paragraph('Construction Contractor: Midwest Renewables Construction')
doc2.add_paragraph('Power Purchaser: MidAmerican Energy')
doc2.save('test_files/wind_project_specs.docx')

# Create first PDF - Battery Storage Project
pdf1 = FPDF()
pdf1.add_page()
pdf1.set_font("Arial", "B", 16)
pdf1.cell(200, 10, txt="GRIDSTORE BATTERY PROJECT", ln=1, align='C')
pdf1.set_font("Arial", size=12)

pdf1.cell(200, 10, txt="PROJECT SUMMARY", ln=1, align='L')
pdf1.set_font("Arial", size=10)
pdf1.multi_cell(0, 10, txt="Project Name: GridStore Battery Energy Storage System\nDeveloper: StorageTech Solutions Inc.\nLocation: Riverside County, California\nTechnology: Lithium-Ion Battery Storage\nCapacity: 100 MW / 400 MWh")

pdf1.set_font("Arial", "B", 12)
pdf1.cell(200, 10, txt="KEY DATES", ln=1, align='L')
pdf1.set_font("Arial", size=10)
pdf1.multi_cell(0, 10, txt="Agreement Date: March 1, 2024\nConstruction Start: September 2024\nCommercial Operation: March 2025")

pdf1.set_font("Arial", "B", 12)
pdf1.cell(200, 10, txt="PROJECT PARTICIPANTS", ln=1, align='L')
pdf1.set_font("Arial", size=10)
pdf1.multi_cell(0, 10, txt="Developer: StorageTech Solutions Inc.\nEPC Contractor: PowerBuild Systems\nBattery Supplier: LG Energy Solution\nOfftaker: California ISO")
pdf1.output("test_files/battery_storage_project.pdf")

# Create second PDF - Solar Plus Storage
pdf2 = FPDF()
pdf2.add_page()
pdf2.set_font("Arial", "B", 16)
pdf2.cell(200, 10, txt="SUNSTORE HYBRID PROJECT", ln=1, align='C')
pdf2.set_font("Arial", size=12)

pdf2.cell(200, 10, txt="HYBRID POWER PROJECT DETAILS", ln=1, align='L')
pdf2.set_font("Arial", size=10)
pdf2.multi_cell(0, 10, txt="Project Name: SunStore Hybrid Facility\nDeveloper: HybridPower Renewables\nLocation: Clark County, Nevada\nSolar Capacity: 200 MW (AC)\nBattery Capacity: 100 MW / 300 MWh")

pdf2.set_font("Arial", "B", 12)
pdf2.cell(200, 10, txt="DEVELOPMENT TIMELINE", ln=1, align='L')
pdf2.set_font("Arial", size=10)
pdf2.multi_cell(0, 10, txt="PPA Execution: January 15, 2024\nConstruction Start: October 1, 2024\nPhase 1 (Solar) COD: September 30, 2025\nPhase 2 (Battery) COD: December 31, 2025")

pdf2.set_font("Arial", "B", 12)
pdf2.cell(200, 10, txt="KEY PARTICIPANTS", ln=1, align='L')
pdf2.set_font("Arial", size=10)
pdf2.multi_cell(0, 10, txt="Lead Developer: HybridPower Renewables\nEPC Partner: RenewBuild Construction\nPanel Supplier: First Solar\nBattery Provider: Tesla\nOfftaker: Nevada Power Company")
pdf2.output("test_files/hybrid_project.pdf") 